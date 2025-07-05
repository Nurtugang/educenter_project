from django.urls import reverse
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required

import io
import docx
import json
import random
from PIL import Image
from docx import Document
from datetime import datetime
from docx.shared import RGBColor
from .models import Question, AnswerOption, UserAnswer, QuizSubject, get_lang_choices


def quizer_choose(request):
	qs = QuizSubject.objects.all()
	context = {
		'qs': qs,
	}
	return render(request, 'quizer_choose.html', context)


def count_correct_answers(user_answers):
	correct_count = 0
	for user_answer in user_answers:
		if user_answer.answer.is_correct:
			correct_count += 1	
	return correct_count


@login_required
def generate_random_test(request, quiz_id):
	doc = Document()
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = docx.shared.Pt(12)

	header = doc.sections[0].header
	paragraph = header.paragraphs[0]
	run = paragraph.add_run('senim-school.com\n')

	current_date = datetime.now().__str__()
	run.add_text(current_date)
	run.add_break()
	paragraph.alignment = 2

	table = doc.add_table(rows=1, cols=2)
	table.columns[0].width = docx.shared.Inches(4.5)
	table.columns[1].width = docx.shared.Inches(2)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = ''  # Эта ячейка остается пустой для выравнивания
	hdr_cells[1].text = 'ФИО:___________________________\n'
	for cell in hdr_cells:
		cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

		
	h = doc.add_paragraph()
	h.alignment = 1
	h.add_run('Тест').bold = True
	h.style.font.name = 'Times New Roman'
	h.style.font.size = docx.shared.Pt(16)

	idx = 1
	questions = Question.objects.filter(quiz_subject__id=quiz_id).prefetch_related('options').order_by('?')[:20]
	shuffle_options = QuizSubject.objects.get(id=quiz_id).shuffle_options
	for question in questions:

		options = list(question.options.all())
		if shuffle_options:
			random.shuffle(options)
		
		p = doc.add_paragraph()
		p.add_run(f'Вопрос №{idx}: {question.text}').bold = True
		p.style.font.name = 'Times New Roman'
		p.style.font.size = docx.shared.Pt(12)
		
		if question.image:
			img = Image.open(question.image.path)
			img_width, img_height = img.size
			page_width_pixels = 900
			if img_width < page_width_pixels:
				scale_factor = page_width_pixels / img_width
				new_width = int(img_width * scale_factor)
				new_height = int(img_height * scale_factor)
				img = img.resize((new_width, new_height))
			img_stream = io.BytesIO()
			img.save(img_stream, format='PNG')  # Change format if necessary
			img_stream.seek(0)
			img_paragraph = doc.add_paragraph()
			img_run = img_paragraph.add_run()
			img_run.add_picture(img_stream, width=docx.shared.Inches(6.5))

		options_paragraph = doc.add_paragraph()
		for i, opt in enumerate(options):
			letter = chr(65+i)
			option_text = f'{letter}) {opt.text}\n'
			if opt.is_correct:  # Check if the option is correct
				is_bold = True
				
			else:
				is_bold = False
			
			option_text = option_text[:-1]  # Remove the trailing newline character
			options_paragraph.add_run(option_text).bold = is_bold  # Bold the correct option
			options_paragraph.add_run('\n')  # Add the newline character back
		idx += 1

	file_stream = io.BytesIO()
	doc.save(file_stream)
	file_stream.seek(0)

	response = HttpResponse(file_stream.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
	response['Content-Disposition'] = f'attachment; filename="Generated Quiz{quiz_id}({request.user.username}).docx"'
	return response


def generate_document(request, quiz_id, attempt_number):
	doc = Document()
	doc.add_heading('Результаты теста', level=1)
	doc.add_paragraph(f'Пользователь: {request.user.first_name} {request.user.last_name}')
	
	user_answers = UserAnswer.objects.filter(user=request.user, question__quiz_subject__id=quiz_id, attempt_number=attempt_number)
	doc.add_paragraph(f'Общее количество вопросов: {len(user_answers)}')
	doc.add_paragraph(f'Количество правильных ответов: {count_correct_answers(user_answers)}')
	
	idx = 1
	for ua in user_answers:
		question = ua.question
		correct_answer = [opt for opt in question.options.all() if opt.is_correct][0]

		doc.add_heading(f'Вопрос №{idx}: {question.text}', level=2)
		if question.image:
			img = Image.open(question.image.path)
			img_width, img_height = img.size
			page_width_pixels = 900
			if img_width < page_width_pixels:
				scale_factor = page_width_pixels / img_width
				new_width = int(img_width * scale_factor)
				new_height = int(img_height * scale_factor)
				img = img.resize((new_width, new_height))
			
			img_stream = io.BytesIO()
			
			img.save(img_stream, format='PNG') 
			img_stream.seek(0)
			doc.add_picture(img_stream, width=docx.shared.Inches(6.5))

		options_paragraph = doc.add_paragraph('Варианты:\n')
		for opt in question.options.all():
			option_text = f'- {opt.text}\n'
			if ua.answer == opt:  # If this option is the user's answer
				if ua.answer == correct_answer:  # If user's answer is correct
					color = RGBColor(0, 128, 0)  # Green
				else:  # If user's answer is wrong
					color = RGBColor(255, 0, 0)  # Red
			elif correct_answer == opt:
				color = RGBColor(0, 128, 0)  # Red
			else:  # If this option is not the user's answer
				color = RGBColor(0, 0, 0)
			options_paragraph.add_run(option_text).font.color.rgb = color  # Red

		idx += 1
	file_stream = io.BytesIO()
	doc.save(file_stream)
	file_stream.seek(0)
	
	response = HttpResponse(file_stream.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
	response['Content-Disposition'] = f'attachment; filename="Quiz Results({request.user.username}).docx"'
	return response

@login_required
def quiz_page(request):
	quiz_id = str(request.GET.get('quiz_id'))
	is_student = request.user.groups.filter(name__in=['Студент']).exists()
	subject = QuizSubject.objects.get(id=quiz_id)
	language = QuizSubject.objects.get(id=quiz_id).language
	time_limit_minutes = QuizSubject.objects.get(id=quiz_id).time_limit_minutes
	questions = Question.objects.prefetch_related('options').filter(quiz_subject__name=subject.name, quiz_subject__language=language).order_by('?')[:20]
	questions_data = []
	for question in questions:
		options = list(question.options.values('text', 'is_correct'))
		questions_data.append({
			'text': question.text,
			'options': options
		})
	questions_json = json.dumps(questions_data)
	context = {
		'questions': questions,
		'questions_json': questions_json,
		'is_student': is_student,
		'subject': subject.name,
		'language': get_lang_choices(language),
		'quiz_id': quiz_id,
		'time_limit_minutes': time_limit_minutes,
		'shuffle_options': subject.shuffle_options
	}
	return render(request, 'quiz_page.html', context)


def testing_handle(request, quiz_id):
	user_answer = UserAnswer.objects.filter(user=request.user, question__quiz_subject__id=quiz_id).order_by('-id').first()
	if user_answer is None:
		last_attempt_id = 0
	else:
		last_attempt_id = user_answer.attempt_number

	if request.method == 'POST':
		for key, value in request.POST.items():
			if key.startswith('question_'):
				question_id = key.split('_')[1]
				selected_option_id = value
				UserAnswer.objects.create(user=request.user, question=Question.objects.get(id=question_id), answer=AnswerOption.objects.get(id=selected_option_id),
							  attempt_number=last_attempt_id+1)
				
		target_url = reverse('quiz_results', args=[quiz_id])
		return redirect(target_url)
	else:
		# Handle non-POST request
		return redirect('quizer_choose')


def quiz_results(request, quiz_id):
	user_attempts = UserAnswer.objects.filter(user=request.user, question__quiz_subject_id=quiz_id).values_list('attempt_number', flat=True).distinct().order_by('-attempt_number')
	results = []
	
	# Данные для статистики
	all_scores = []
	attempt_data = []
	
	for attempt_number in user_attempts:
		# Get user answers for the current attempt
		user_answers = UserAnswer.objects.filter(user=request.user, question__quiz_subject_id=quiz_id, attempt_number=attempt_number)
		total_questions_count = user_answers.count()

		# Выбрать только неправильно решенные тестовые задания
		wrong_answers = user_answers.exclude(answer__is_correct=True)
		wrong_answers_count = len(wrong_answers)
		correct_answers_count = total_questions_count - wrong_answers_count
		
		# Процент правильных ответов
		percentage = (correct_answers_count / total_questions_count * 100) if total_questions_count > 0 else 0
		all_scores.append(percentage)
		
		# Данные для графика
		attempt_data.append({
			'attempt': attempt_number,
			'score': percentage,
			'correct': correct_answers_count,
			'total': total_questions_count
		})
		
		# Соответстующие правильные ответы, неправильно решенных тестовых заданий
		wrong_questions = [answer.question for answer in wrong_answers]

		# Step 2: Get the correct answers for the wrong questions
		correct_answers = []
		for question in wrong_questions:
			correct_answer = AnswerOption.objects.filter(question=question, is_correct=True).first()
			correct_answers.append(correct_answer)

		results.append({
			'attempt_number': attempt_number,
			'correct_answers_count': correct_answers_count,
			'total_questions_count': total_questions_count,
			'percentage': round(percentage, 1),
			'data': zip(wrong_answers, correct_answers)
		})

	# Вычисляем статистику
	stats = {}
	if all_scores:
		stats = {
			'total_attempts': len(all_scores),
			'average_score': round(sum(all_scores) / len(all_scores), 1),
			'best_score': round(max(all_scores), 1),
			'last_score': round(all_scores[0], 1) if all_scores else 0,  # первый в списке = последний по времени
		}
		
		# Сортируем attempt_data по номеру попытки для правильного отображения на графике
		attempt_data.sort(key=lambda x: x['attempt'])
	else:
		stats = {
			'total_attempts': 0,
			'average_score': 0,
			'best_score': 0,
			'last_score': 0,
		}

	# Добавляем изменение относительно предыдущей попытки
	for i, result in enumerate(results):
		if i < len(results) - 1:  # если не последний элемент (не первая попытка)
			prev_percentage = results[i + 1]['percentage']  # предыдущая попытка
			current_percentage = result['percentage']
			change = current_percentage - prev_percentage
			result['change'] = round(change, 1)
			result['change_direction'] = 'up' if change > 0 else 'down' if change < 0 else 'same'
		else:
			result['change'] = None
			result['change_direction'] = 'first'

	# Подготавливаем данные для графиков в JSON формате
	chart_data = {
		'attempts': [item['attempt'] for item in attempt_data],
		'scores': [item['score'] for item in attempt_data],
		'detailed_data': attempt_data
	}

	context = {
		'results': results, 
		'quiz_id': quiz_id,
		'quiz': QuizSubject.objects.get(id=quiz_id).name,
		'language': get_lang_choices(QuizSubject.objects.get(id=quiz_id).language),
		'stats': stats,
		'chart_data': json.dumps(chart_data)
	}
	return render(request, 'quiz_results.html', context)